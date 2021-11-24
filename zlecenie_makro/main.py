from json.decoder import JSONDecodeError
from os import error, path
import requests
from bs4 import BeautifulSoup
import time
import json
import pandas as pd
import traceback
from pdfminer.high_level import extract_text
import pdfminer
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import mammoth

#uzyskaj dane do logowania
try:
    with open('makro_user.txt') as user_cfg:
        USERNAME = user_cfg.readline()
        PWD = user_cfg.readline()
except:
    print('BŁĄD PLIKU ZAWIERAJĄCEGO DANE DO LOGOWANIA')
    print(traceback.format_exc())


# produkt i jego atrybuty
class Product():
    def __init__(self, index, ean, product_name, photo_url, net_price, gross_price, jm, brand, manufacturer, origin, vat, group=None, subgroup=None, subsubgroup=None, net_weight=None, gross_weight=None, attachment=False):
        self.index = index
        self.ean = ean
        self.product_name = product_name
        self.photo_url = photo_url
        self.net_price = net_price
        self.gross_price = gross_price
        self.jm = jm
        self.brand = brand
        self.manufacturer = manufacturer
        self.origin = origin
        self.vat = vat
        self.net_weight = net_weight
        self.gross_weight = gross_weight
        self.attachment = attachment
        self.group = group
        self.subgroup = subgroup
        self.subsubgroup = subsubgroup

    def as_dict(self):
        return {
            'index': [self.index], 
            'ean': [self.ean],
            'product_name': [self.product_name],
            'photo_url': [self.photo_url],
            'net_price': [self.net_price],
            'gross_price': [self.gross_price],
            'jm': [self.jm],
            'brand': [self.brand],
            'manufacturer': [self.manufacturer],
            'origin': [self.origin],
            'vat': [self.vat],
            'net_weight': [self.net_weight],
            'gross_weight': [self.gross_weight],
            'group': [self.group],
            'subgroup': [self.subgroup],
            'subsubgroup': [self.subsubgroup],
            'attachment': [self.attachment],
            }

class Processed_OCR():
    def __init__(self, prd_id, add_descr='', ean='', ingridients='', allergens='', before_table='', table1=[], table2=[], table3=[], additional_info={}):
        self.prd_id = prd_id
        self.add_descr = add_descr
        self.ean = ean
        self.ingridents = ingridients
        self.allergens = allergens
        self.before_table = before_table
        self.table1 = table1
        self.table2 = table2
        self.table3 = table3
        self.additional_info = additional_info

def get_bearer_token():
    # Zaloguj użytkownika, uzyskaj token potrzebny do logowania oraz uzupełnij brakujące nagłówki
    headers = {
        'Accept-Encoding':'gzip, deflate, br',
        'Connection':'keep-alive',
        'sec-ch-ua':'"Chromium";v="94", "Google Chrome";v="94", ";Not A Brand";v="99"',
        'sec-ch-ua-mobile':'?0',
        'sec-ch-ua-platform':'"Windows"',
        'accept':'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
        'sec-fetch-site':'cross-site', 
        'sec-fetch-user':'?1', 
        'sec-fetch-mode':'navigate',
        'Sec-Fetch-Dest':'document',
        'upgrade-insecure-requests':'1', 
        'User-Agent' : 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/94.0.4606.81 Safari/537.36',
        'Cache-Control':'max-age=0',
        'accept-language':'en-GB,en;q=0.9',
        'referer':'https://makro.ehurtownia.pl/'
    }

    main_page_url = 'https://sso.infinite.pl/auth/realms/InfiniteEH/protocol/openid-connect/auth?client_id=ehurtownia-panel-frontend&redirect_uri=https%3A%2F%2Fmakro.ehurtownia.pl%2Fone%2F%3Fredirect_fragment%3D%252F&response_mode=fragment&response_type=code&scope=openid'

    session = requests.Session()

    while True:
        try:
            print('Otwieram stronę z logowaniem...')
            main_page = session.get(main_page_url, headers=headers)
            print('Strona została pomyślnie otwarta.\n')
            print('Wyszukuję link do logowania...')
            soup = BeautifulSoup(main_page.text, 'html.parser')
            form = soup.find('form', attrs={'id':"kc-form-login"})
            login_url = form['action']
            print('Link został znaleziony.\n')

            login_data = {
                'username': USERNAME,
                'password': PWD,
                'credentialId':''
            }

            print('Próbuję zalogować użytkownika...')
            auth_user = session.post(login_url, headers=headers, data=login_data, allow_redirects=True)
            
            try:
                code_index = (auth_user.url).find('code') + len('code') + 1
                print('Użytkownik zalogowany pomyślnie!\n')
                print('Uzyskuję kod autoryzacyjny...')
                auth_code = (auth_user.url)[code_index:]
                print('Kod autoryzacyjny uzyskany.\n')
            except:
                print('Niepoprawne dane logowania!')
                continue    

            auth_data = {
                'code': auth_code,
                'grant_type': 'authorization_code',
                'client_id': 'ehurtownia-panel-frontend',
                'redirect_uri': 'https://makro.ehurtownia.pl/one/?redirect_fragment=%2F'
            }

            print('Uzyskuję token autoryzacyjny...')
            token_page = requests.post('https://sso.infinite.pl/auth/realms/InfiniteEH/protocol/openid-connect/token', headers=headers, data=auth_data)

            resp = json.loads(token_page.text)
            acc_token = resp['access_token']
            headers['Authorization'] = 'Bearer '+ acc_token
            print('Token autoryzacyjny został pomyślnie uzyskany!\n')
            return session, headers

        except:
            print('Wystąpił problem z uzyskaniem tokenu!')
            print('Zapisuję błąd...')
            with open('error_log.txt', 'a') as error_log:
                err = traceback.format_exc()
                error_log.write(err)
            time.sleep(3)
            print('Próbuję ponownie...')
            continue

def get_prd_data(session, headers, query_index, query_ean, unavailable_products):
    # Poprzez API uzyskaj dane produktów - część bezpośrednio poprzez wyszukiwarkę, resztę poprzez kartę produktu, pobierz załącznik.
    # Jeżeli produkt jest niedostepny zapisz go na liście niedostępnych i akcje zapisz w logu.
    # Jeżeli sesja wygasła, zaloguj użytkownika ponownie
    product_page_url = 'https://makro.ehurtownia.pl/one/eh-one-backend/rest/1053/8/110100/oferta/%s/karta-towaru?lang=PL'
    query_url = 'https://makro.ehurtownia.pl/one/eh-one-backend/rest/1053/8/110100/oferta?lang=PL&offset=0&limit=10&sortAsc=nazwa&cechaWartosc=%s %s'
    attachment_url = 'https://makro.ehurtownia.pl/one/eh-one-backend/rest/1053/8/110100/oferta/%s/zalacznik?lang=PL'

    attempt = 0
    product_name = ''
    net_weight = ''
    while attempt < 4:
        try:
            print(f"Uzyskuję dane dla towaru z indekstem {query_index} i kodem EAN {query_ean }")
            query_result = session.get((query_url % (query_index, query_ean )), headers=headers, timeout=7)
            product_details = json.loads(query_result.text)
            try:
                index = query_index
                ean = query_ean
                product_details = product_details['pozycje'][0]
                product_name =  product_details['nazwa']
                photo_url = product_details['zdjecie']
                try:
                    net_price = product_details['cenaNettoJedn']
                    gross_price = product_details['cenaBruttoJedn']
                except:
                    net_price = product_details['cenaNetto']
                    gross_price = product_details['cenaBrutto']
                jm = product_details['jm']
                brand = product_details['marka']
                manufacturer = product_details['producenciNazwa']
                origin = product_details['krajPochodzeniaNazwa']
                vat = product_details['procVat']
                product = Product(index, ean, product_name, photo_url, net_price, gross_price, jm, brand, manufacturer, origin, vat)
                print(f'Podstawowe dane dla towaru z indekstem {query_index} ({product_name} uzyskane!')
            except IndexError:
                print('Produkt nie jest dostępny! Dodaję go do listy niedostępnych produktów')
                unavailable_products['Indeks/SupplerItemCode'].append(query_index)
                unavailable_products['Kod kreskowy/EAN'].append(query_ean)
                with open('log.txt', 'a') as log:
                    log.write(f'Produkt niedostępny: Indeks {query_index}, EAN {query_ean}\n')
                return session, headers, None, unavailable_products
            else:
                product_page = session.get((product_page_url % query_index), headers=headers, timeout=7)
                product_details = json.loads(product_page.text)
                try:
                    net_weight = product_details['wagaNettoJm']
                    gross_weight = product_details['wagaBruttoJm']
                    group = product_details['eGrupyNazwa']
                    subgroup = product_details['ePodGrupyNazwa']
                    subsubgroup = product_details['kluczHierGrupTowarowych']['ePodPodGrupyKod']
                    product.net_weight = net_weight
                    product.gross_weight = gross_weight
                    product.group = group
                    product.subgroup = subgroup
                    product.subsubgroup = subsubgroup
                    print(f'Dodatkowe dane dla towaru z indekstem {query_index} ({product_name} uzyskane!')
                except:
                    print('Dodatkowe info o produkcie o EAN: %s INDEX: %s jest niedostępne' % (query_index, query_ean))
                    with open('log.txt', 'a') as log:
                        log.write(f'Dodatkowe informacje (waga netto, waga brutto) niedostępne dla produktu: Indeks {query_index}, EAN {query_ean}\n')
                else:
                    try:
                        pdf_id = product_details['zalaczniki'][0]['id']
                        attachment_page = session.get((attachment_url % pdf_id), headers=headers, timeout=7)
                        if attachment_page.status_code == 200:
                            file_name = 'zalaczniki_pdf\\' + str(query_ean) + '.pdf'
                            with open(file_name, 'wb') as pdf_file:
                                pdf_file.write(attachment_page.content)
                                product.attachment = True
                                print(f'Załącznik dla towaru z indekstem {query_index} ({product_name} uzyskany!')
                    except:
                        with open('log.txt', 'a') as log:
                            log.write(f'Brak załącznika dla produktu: Indeks {query_index}, EAN {query_ean}\n')
        except (ConnectionResetError, NameError,requests.exceptions.Timeout,requests.exceptions.ConnectionError, TimeoutError, requests.exceptions.ChunkedEncodingError):
            print('Błąd ładowania strony! Próbuję ponownie...')
            attempt += 1
        except:
            if "INTERNAL_ERROR" in query_result.text or "Unable to authenticate bearer token" in query_result.text and attempt < 3:
                session, headers = get_bearer_token()
                attempt += 1
            else:
                print('Zapisuję błąd do pliku "error_log.txt"...') 
                with open('error_log.txt', 'a') as error_log:
                    err = traceback.format_exc()
                    error_log.write(err)
                with open('log.txt', 'a') as log:
                    log.write(f'Wystąpił błąd dla produktu: Indeks {query_index}, EAN {query_ean}\n')
                return session, headers, None, unavailable_products
        else:
            return session, headers, product, unavailable_products
            
    return session, headers, None, unavailable_products

def has_numbers(string):
    # Sprawdź czy w tekście znajduje/ą się cyfry
    return any(char.isdigit() for char in string)

def open_pdf(file_name):
    # Otwórz plik pdf (załączniks), podziel zawartość do listy, zamień elementy które nie są poprawnie odczytywane /
    # oraz usuń puste elementy (spacje)
    try:
        text = extract_text(file_name)
        text = text.replace('½', '1/2')
        text = text.replace('¼', '1/4')
        text = text.replace('¾', '3/4')
        text = text.replace('ø', 'średnica')
        text = text.replace('≥', '>=')
        text = text.replace('º', 'o')
        text = text.replace('≤', '<=')
        text = text.replace('ù', 'u')
        text = text.replace('è', 'e')
        text = text.replace('≈', '~')
        text = text.replace('α', 'a')
        text = text.replace('ñ', 'n')
        text = text.replace('à', 'a')
        text = text.replace('¯', '-')
        text = text.replace('ˉ', '-')

        split_text = text.splitlines()
        split_text = [i for i in split_text if i]
        return split_text
    except pdfminer.pdfparser.PDFSyntaxError:
        print('Błąd PDF!')
        return None

def rem_prd_id(split_text, prd_id):
    # Usuń elementy zawierające nazwe produktu
    el_to_rem = []
    prd_id = prd_id.lower()
    starting_prd_id = prd_id
    for i, line in enumerate(split_text):
        line = line.lower()
        if line in prd_id:
            if prd_id == starting_prd_id and line[0] != prd_id[0]:
                pass 
            else:
                el_to_rem.append(i)
                prd_id = prd_id.replace(line,'')
                if prd_id == '':
                    break

    for i in reversed(el_to_rem):
        split_text.pop(i)

    return split_text

def get_ean_allergens_add_discr(split_text):
    # Wyciągnij z tekstu alergeny (jeżeli istnieją w tej części) oraz opis dodatkowy (pod nazwą)
    # Usuń z tekstu uzyskane elementy
    i = 0
    add_descr = ''
    allergens = ''
    el_to_rem = []
    while 'kJ' not in split_text[i] and 'kcal' not in split_text[i+1]:
        if "GTIN" in split_text[i]:
            i+=1
            break
        if "Contains :" in split_text[i] or allergens != '':
            allergens += split_text[i]
        else:
            add_descr += split_text[i] + ''
        el_to_rem.append(i)
        i+=1

    allergens = allergens.replace('Contains : ', '')
    if allergens != '':
        try:
            split_text.remove('Obecność alergenu')
        except ValueError:
            pass


    for i in reversed(el_to_rem):
        split_text.pop(i)

    return split_text, allergens, add_descr

def rem_ean(split_text, ean):
    # Usuń z tekstu część zawierającą kod EAN
    ean_part = 'GTIN / EAN : ' + ean
    try:
        split_text.remove(ean_part)
    except ValueError:
        for el in split_text:
            if 'GTIN / EAN' in el:
                split_text.remove(el)
                return split_text, ean
    return split_text, ean

def get_ingridients(split_text):
    # Uzyskaj listę składników
    # Usuń z tekstu uzyskane elementy
    possible_end = ("Referencyjna Wartość Spożycia", "Sugerowana porcja produktu (wartość)", "Warunki przechowywania", "Opakowanie oznakowane datą minimalnej trwałości", \
                        "Obowiązkowe dodatkowe informacje wymagane na etykiecie", "Kraj pochodzenia", "Sposób przygotowania", "Propozycja podania", "Przybliżona liczba porcji", \
                        "Zawartość netto", "Pozostałe informacje", "Nazwa opisowa produktu", "Nazwa firmy","Adres", "Forma kontaktu", "Inne informacje dla konsumenta"\
                        "Obecność alergenu", "Wartość odżywcza/ 100g", "kJ", "kcal", "Contains", "Obecność alergenu")
    el_to_rem = []
    try:
        index_ingr = split_text.index("Składnik") + 1
        el_to_rem.append(index_ingr-1)
        if index_ingr != 0:
            ingridients = ''
            while all(x not in split_text[index_ingr] for x in possible_end):
                if split_text[index_ingr] != 'Obecność alergenu':
                    if "SKŁADNIKI" in split_text[index_ingr]:
                        ingridients += split_text[index_ingr].replace('SKŁADNIKI', '')
                    else:
                        ingridients += split_text[index_ingr]
                el_to_rem.append(index_ingr)
                index_ingr += 1
            ingridients = ingridients.split()
        else:
            ingridients = None

        for i in reversed(el_to_rem):
            split_text.pop(i)
    except ValueError:
        return split_text, []

    return split_text, ingridients

# def get_words_to_bold(ingridients_split):
#     # Wybierz słowa do pogrubienia
#     for word in ingridients_split:
#         if word.isupper() and not has_numbers(word):
#             for char in word:
#                 # do not bold any of these characters
#                 if char in ('(', ')', ',', ':'):
#                     #print(char)
#                     pass
#                 else:
#                     # bold char
#                     #print(char)
#                     pass

def get_allergens(split_text, allergens):
    # Uzyskaj alergeny jeżeli nie zostały już wcześniej znalezione
    # Usuń z tekstu uzyskane elementy
    started = True

    if allergens == '':
        el_to_rem = []
        try:
            index_all = split_text.index('Obecność alergenu') + 1  
            el_to_rem.append(index_all-1) 
            if 'Wartość odżywcza/ 100g' not in split_text[index_all]:
                if index_all != 0:
                    allergens = split_text[index_all]
                    allergens = allergens.replace('Contains : ','')
                    el_to_rem.append(index_all) 
        except ValueError:
            allergens = None

    try:
        if split_text[index_all] == 'Wartość odżywcza/ 100g':
            started = False
        elif allergens == None:
            started = False           
    except UnboundLocalError:
        if allergens == None:
            started = False
        
    if started == False:
        for i, el in enumerate(split_text):
            if 'Contains' in el and not started:
                el = el.replace('Contains : ','')
                allergens = el
                el_to_rem.append(i)
                started = True
            elif started and all(x not in el for x in ('Obecność alergenu', 'Wartość odżywcza/ 100g', 'Wartość odżywcza/ 100g','kJ','kcal')):
                allergens += el
                el_to_rem.append(i)
            elif started:
                break

    try:
        for i in reversed(el_to_rem):
                split_text.pop(i)
    except UnboundLocalError:
        pass

    return split_text, allergens

def get_additional_info(split_text):
    # Uzyskaj dodatkowe informacje znajdujące się na samym końcu załącznika. 
    # Format = klucz : wartość
    additional_info = {
        "Referencyjna Wartość Spożycia": "",
        "Sugerowana porcja produktu (wartość)": "",
        "Warunki przechowywania": "",
        "Opakowanie oznakowane datą minimalnej trwałości": "",
        "Obowiązkowe dodatkowe informacje wymagane na etykiecie": "",
        "Inne informacje dla konsumenta": "",
        "Kraj pochodzenia": "",
        "Sposób przygotowania": "",
        "Propozycja podania": "",
        "Przybliżona liczba porcji": "",
        "Zawartość netto": "",
        "Pozostałe informacje": "",
        "Nazwa opisowa produktu": "",
        "Forma kontaktu": "",
        "Nazwa firmy": "",
        "Adres": "",
    }

    additional_info_rows = ["Referencyjna Wartość Spożycia", "Sugerowana porcja produktu (wartość)", "Warunki przechowywania", "Opakowanie oznakowane datą minimalnej trwałości", \
                            "Obowiązkowe dodatkowe informacje wymagane na etykiecie", "Inne informacje dla konsumenta", "Kraj pochodzenia", "Sposób przygotowania", "Propozycja podania", "Przybliżona liczba porcji", \
                            "Zawartość netto", "Pozostałe informacje", "Nazwa opisowa produktu", "Forma kontaktu", "Nazwa firmy", "Adres"]

    el_to_rem = []
    for key, val in additional_info.items():
        val = ""
        el_found = False
        for el in split_text:
            if key in el:
                el_found = True
                val += el[(len(key)+3):]
                el_to_rem.append(split_text.index(el))
            elif el_found:
                if any(x in el for x in additional_info_rows):
                    break
                val += el
                el_to_rem.append(split_text.index(el))
        additional_info[key] = val
    
    el_to_rem = list(set(el_to_rem))
    el_to_rem.sort()
    for i in reversed(el_to_rem):
        split_text.pop(i)
    
    return split_text, additional_info

def get_tables(split_text):
    # Uzyskaj tabele : jeżeli jest pusta - pomiń ją
    # possible_rows - możliwe wiersze w tabeli
    # jump - zmienna służąca do zmiany odległości między nazwami wierszy a wartosciami

    possible_rows = ('energia', 'wartość energetyczna', 'tłuszcz', 'w tym kwasy tłuszczowe nasycone', 'węglowodany', 'w tym cukry', \
                'białko', 'sól', 'niacyna', 'witamina B6', 'kwas pantotenowy', 'witamina', 'kwas', 'ryboflawina', 'biotyna', 'wapń',\
                'tiamina', 'cukry', 'kwasy tłuszczowe nasycone', 'sód', 'kwasy', 'błonnik', 'składniki mineralne', 'minerały', 'fosfor',\
                'żelazo', 'magnez', 'cynk', 'wartość odżywcza/ 100g')

    before_table_arr = [x for x in split_text if x.lower() in ('w 100 g', 'na 100 g', 'w 100 ml', 'na 100ml', 'na 100g', 'w 100g', 'na 100 ml', 'w 100ml')]
    if before_table_arr != []:
        jump = 1
        before_table = before_table_arr[0]
    else:
        before_table = ''
        jump = 0

    table_counter = 0
    for i in range(len(split_text) - 1):
        if ('kJ' in split_text[i] and has_numbers(split_text[i])) and ('kcal' in split_text[i+1] and has_numbers(split_text[i+1])):
            table_counter += 1
        

    nutrinional_val_arr_1 = []
    nutrinional_val_arr_2 = []
    nutrinional_val_arr_3 = []

    if table_counter == 1:
        try:
            arr_indx = split_text.index('Wartość odżywcza/ 100g') + 2 + jump
        except ValueError:
            return split_text, before_table, nutrinional_val_arr_1, nutrinional_val_arr_2, nutrinional_val_arr_3
        else:
            starting_index = arr_indx
            rows_counter = 0
            while any(x in split_text[arr_indx].lower() for x in possible_rows) and all(x[0] != split_text[arr_indx] for x in nutrinional_val_arr_1[1:]):
                nutrinional_val_arr_1.append([split_text[arr_indx], ''])
                arr_indx += 1
                rows_counter += 1
            for i, x in enumerate(range(rows_counter, 0, -1)):
                nutrinional_val_arr_1[i][1] = split_text[starting_index - 2 - jump - x]

            if nutrinional_val_arr_1[0][0] == 'Wartość odżywcza/ 100g':
                return split_text, '', [], [], []

            try:
                start_second_try = False
                if nutrinional_val_arr_1 == []:
                    start_second_try = True
                elif nutrinional_val_arr_1[0][1] and nutrinional_val_arr_1[0][0]:
                    if has_numbers(nutrinional_val_arr_1[0][1]) and nutrinional_val_arr_1[0][0] == 'Wartość energetyczna' and 'kJ' not in nutrinional_val_arr_1[0][1] and "%" not in nutrinional_val_arr_1[0][1]:
                        start_second_try = True
            except IndexError:
                start_second_try = True

            if start_second_try == True:
                try:
                    old_starting_indx = starting_index  
                    nutrinional_val_arr_1 = []
                    try:
                        arr_indx = split_text.index('wartość odżywcza/ porcja') + 2 + jump
                    except ValueError:
                        return split_text, before_table, nutrinional_val_arr_1, nutrinional_val_arr_2, nutrinional_val_arr_3
                    else:
                        starting_index = arr_indx
                        rows_counter = 0
                        while any(x in split_text[arr_indx].lower() for x in possible_rows) and all(x[0] != split_text[arr_indx] for x in nutrinional_val_arr_1[1:]):
                            nutrinional_val_arr_1.append([split_text[arr_indx], ''])
                            arr_indx += 1
                            rows_counter += 1
                        for i, x in enumerate(range(rows_counter, 0, -1)):
                            nutrinional_val_arr_1[i][1] = split_text[old_starting_indx - 2 - jump - x]
                    try:
                        if nutrinional_val_arr_1[0][1]:
                            for i in range((len(nutrinional_val_arr_1) - 1)):
                                if nutrinional_val_arr_1[i][1] == '- kJ' and nutrinional_val_arr_1[i+1][1] == '- kcal':
                                    return split_text, '', [], [], []
                    except IndexError:
                        pass
                except:
                    with open('error_log.txt', 'a') as error_log:
                        err = traceback.format_exc()
                        error_log.write(err)                

    elif table_counter == 2:
        try:
            arr_indx = split_text.index('wartość odżywcza/ porcja') + 2 + jump
            starting_index = arr_indx
            rows_counter = 0
            while any(x in split_text[arr_indx].lower() for x in possible_rows) and all(x[0] != split_text[arr_indx] for x in nutrinional_val_arr_1[1:]):
                nutrinional_val_arr_1.append([split_text[arr_indx], ''])
                nutrinional_val_arr_2.append([split_text[arr_indx], ''])
                arr_indx += 1
                rows_counter += 1
            arr_indx += 1 #pass "Wartość odżywcza/ GDA"
            for i, val in enumerate(split_text[arr_indx:(arr_indx+rows_counter)]):
                nutrinional_val_arr_1[i][1] = val
            
            for i, val in enumerate(split_text[(arr_indx+rows_counter):(arr_indx+2*rows_counter)]):
                nutrinional_val_arr_2[i][1] = val


            start = False
            try:
                if 'kJ' not in nutrinional_val_arr_1[0][1]:
                    start = True
            except IndexError:
                start = True
            finally:
                if start == True:
                    for i, el in enumerate(reversed(split_text)):
                        if i < rows_counter:
                            nutrinional_val_arr_2[rows_counter-1-i][1] = el
                        elif i < 2*rows_counter:
                            nutrinional_val_arr_1[2*rows_counter-1-i][1] = el
        except:
            with open('error_log.txt', 'a') as error_log:
                err = traceback.format_exc()
                error_log.write(err)

        try:
            start_second_try = False
            if nutrinional_val_arr_1 == []:
                start_second_try = True
            elif nutrinional_val_arr_1[0][1] and nutrinional_val_arr_1[0][0]:
                if has_numbers(nutrinional_val_arr_1[0][1]) != True or "%" in nutrinional_val_arr_1[0][1] or (nutrinional_val_arr_1[0][0] == 'Wartość energetyczna' and 'kJ' not in nutrinional_val_arr_1[0][1]):
                    start_second_try = True
        except IndexError:
            start_second_try = True

        if start_second_try == True:
            try:
                start_indx = split_text.index('Wartość odżywcza/ 100g') - 1
                for i in range(rows_counter*2):
                    if i < rows_counter:
                        indx = rows_counter - i - 1
                        nutrinional_val_arr_2[indx][1] = split_text[start_indx-i]
                    elif i < 2*rows_counter:
                        indx = rows_counter - (i - rows_counter) - 1
                        nutrinional_val_arr_1[indx][1] = split_text[start_indx-i]
            except:
                with open('error_log.txt', 'a') as error_log:
                    err = traceback.format_exc()
                    error_log.write(err)

    elif table_counter == 3:
        try:        
            rows_counter = 0
            arr_indx = split_text.index('Wartość odżywcza/ 100g') + 2 + jump
            while any(x in split_text[arr_indx].lower() for x in possible_rows) and all(x[0] != split_text[arr_indx] for x in nutrinional_val_arr_1[1:]):
                nutrinional_val_arr_1.append([split_text[arr_indx], ''])
                nutrinional_val_arr_2.append([split_text[arr_indx], '']) 
                nutrinional_val_arr_3.append([split_text[arr_indx], ''])  
                rows_counter += 1
                arr_indx += 1
            row = 0
            i = 0
            arr_indx += 1
            while arr_indx < (len(split_text) - 1):
                arr_indx += 1
                if has_numbers(split_text[arr_indx]):
                    if i < rows_counter:
                        nutrinional_val_arr_1[row][1] = split_text[arr_indx]
                    elif i < 2*rows_counter:
                        nutrinional_val_arr_2[row][1] = split_text[arr_indx]
                    elif i < 3*rows_counter:
                        nutrinional_val_arr_3[row][1] = split_text[arr_indx]
                    row += 1
                    i += 1
                    if row == rows_counter:
                        row = 0
        except:
            with open('error_log.txt', 'a') as error_log:
                err = traceback.format_exc()
                error_log.write(err)

    else:
        return split_text, '', [], [], []
    #xd
    if nutrinional_val_arr_1:
        cut_txt_el = nutrinional_val_arr_1[0][1]
        cut_txt_indx = split_text.index(cut_txt_el)
        split_text = split_text[:cut_txt_indx]

    return split_text, before_table, nutrinional_val_arr_1, nutrinional_val_arr_2, nutrinional_val_arr_3

def write_to_docx(ocr_data):
    # Wpisz uzyskane dane do dokumentu z rozszerzeniem "docx"

    document = Document()

    prd_id_p1 = document.add_paragraph()
    p1_runner = prd_id_p1.add_run(ocr_data.prd_id)
    p1_runner.bold = True
    p1_font = p1_runner.font
    p1_font.color.rgb = RGBColor(0x1A, 0x3C, 0x7B)
    p1_font.size = Pt(14)
    p1_font.name = "Arial"

    if ocr_data.add_descr != '':
        add_descr_p2 = document.add_paragraph()
        p2_runner = add_descr_p2.add_run(ocr_data.add_descr)
        p2_font = p2_runner.font
        p2_font.size = Pt(8)
        p2_font.name = "Arial"   

    ean_p3 = document.add_paragraph()
    p3_runner = ean_p3.add_run("GTIN / EAN : ")
    p3_runner.bold = True
    p3_runner_2 = ean_p3.add_run(ocr_data.ean)
    p3_font_1 = p3_runner.font
    p3_font_1.size = Pt(8)
    p3_font_1.name = "Arial"
    p3_font_2 = p3_runner_2.font
    p3_font_2.size = Pt(8)
    p3_font_2.name = "Arial"

    if ocr_data.ingridients != '' and ocr_data.ingridients != []:
        ingridients_p4 = document.add_paragraph()
        ingridients_p4.paragraph_format.space_after = Pt(0)
        p4_runner = ingridients_p4.add_run("Składniki")
        p4_runner.bold = True
        p4_font = p4_runner.font
        p4_font.size = Pt(10)
        p4_font.name = "Arial"
        p4_font.color.rgb = RGBColor(0x1A, 0x3C, 0x7B)
        
        ingridients_p5 = document.add_paragraph()
        ingridients_p5.paragraph_format.space_before = Pt(0)
        for word in ocr_data.ingridients:
            p5_runner = ingridients_p5.add_run(word + ' ')
            p5_font = p5_runner.font
            p5_font.size = Pt(8)
            p5_font.name = "Arial"
    
    if ocr_data.allergens != '' and ocr_data.allergens != None:
        allergens_p6 = document.add_paragraph()
        allergens_p6.paragraph_format.space_after = Pt(0)
        p6_runner = allergens_p6.add_run("Obecność alergenów")
        p6_runner.bold = True
        p6_font = p6_runner.font
        p6_font.size = Pt(10)
        p6_font.name = "Arial"
        p6_font.color.rgb = RGBColor(0x1A, 0x3C, 0x7B)
        
        allergens_p7 = document.add_paragraph()
        allergens_p7.paragraph_format.space_before = Pt(0)
        p7_runner = allergens_p7.add_run('Zawiera : ')
        p7_runner.bold = True
        #p7_runner = allergens_p7.add_run('Zawiera : ' + allergens)
        p7_font = p7_runner.font
        p7_font.size = Pt(8)
        p7_font.name = "Arial"     
        p8_runner = allergens_p7.add_run(ocr_data.allergens)
        p8_font = p8_runner.font
        p8_font.size = Pt(8)
        p8_font.name = "Arial"  

    if ocr_data.table3 != []:
        tables = [ocr_data.table1, ocr_data.table2, ocr_data.table3]
        table_ids = ['Wartość odżywcza/ 100g', 'Wartość odżywcza/ porcja', 'Wartość odżywcza/ GDA']
    elif ocr_data.table2 != []:
        tables = [ocr_data.table1, ocr_data.table2]
        table_ids = ['Wartość odżywcza/ 100g', 'Wartość odżywcza/ porcja']
    elif ocr_data.table1 != []:
        tables = [ocr_data.table1]
        table_ids = ['Wartość odżywcza/ 100g']
    else:
        tables = None

    if tables:
        for nutrinional_table, table_id in zip(tables, table_ids):
            table_header_p = document.add_paragraph()
            table_header_p_runner = table_header_p.add_run(table_id)
            table_header_p_runner.bold = True
            header_font = table_header_p_runner.font
            header_font.size = Pt(10)
            header_font.name = "Arial"
            header_font.color.rgb = RGBColor(0x1A, 0x3C, 0x7B)

            if ocr_data.before_table:
                before_table_p = document.add_paragraph()
                before_table_p.paragraph_format.space_after = Pt(0)
                before_table_p_runner = before_table_p.add_run(ocr_data.before_table)
                before_table_p_runner.bold = True
                before_table_font = before_table_p_runner.font
                before_table_font.size = Pt(8)
                before_table_font.name = "Arial"
                before_table_font.color.rgb = RGBColor(0x1A, 0x3C, 0x7B) 

            rows = len(nutrinional_table) + 1
            cols = len(nutrinional_table[0])
            table = document.add_table(rows=rows, cols=cols)
            table.style = 'TableGrid'
            header_row = table.rows[0]
            table.rows[0].cells[0]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="1A3C7B"/>'.format(nsdecls('w'))))
            for i in range(cols):
                a,b = header_row.cells[:2]
                a.merge(b)
            header_row = header_row.cells
            header_row[0].text = table_id
            table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
            table.rows[0].cells[0].paragraphs[0].runs[0].font.name = 'Arial'
            table.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(8)
            table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)

            for i, (id, val) in enumerate(nutrinional_table, 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = id
                row_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
                row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
                row_cells[0].paragraphs[0].runs[0].font.size = Pt(8)
                row_cells[1].text = val    
                row_cells[1].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
                row_cells[1].paragraphs[0].runs[0].font.name = 'Arial'
                row_cells[1].paragraphs[0].runs[0].font.size = Pt(8)


    rand_p = document.add_paragraph()

    for i, (key, val) in enumerate(ocr_data.additional_info.items()):
        if val == '':
            continue
        add_info_p = document.add_paragraph()
        add_info_p.paragraph_format.space_before = Pt(2)
        add_info_p.paragraph_format.space_after = Pt(0)
        p_runner = add_info_p.add_run(key + " : ")
        p_runner.bold = True
        p_runner_2 = add_info_p.add_run(val)
        p = p_runner.font
        p.size = Pt(8)
        p.name = "Arial"
        p_2 = p_runner_2.font
        p_2.size = Pt(8)
        p_2.name = "Arial"     

    docx_file_name = 'zalaczniki_docx\\' + ocr_data.ean + '.docx'
    document.save(docx_file_name)
    return docx_file_name

def write_to_html(file):
    # Przepisz dokument "docx" do dokumentu "html"
    # Używając CSS odtwórz formatowanie dokumentu
    custom_css = """
    <style>
    body{
        font-family: Arial;
        font-size: 14pt;
    }
    
    table{
        border: 1px solid black;
        width: 70%;
        border-collapse:collapse;
    }

    table td{
        border: 1px solid black;
        padding: 3px;
    }

    table strong{
        color: white;
    }

    table td{
        width: 50%;
    }

    table td:nth-child(2){
        text-align:right;
    }
    
    tr:first-child{
        background-color: #1A3C7B;
    }

    p:first-of-type strong{
        font-size: 20pt;
        font-weight: bold;
    }

    strong{
        color: #1A3C7B;
        font-size: 16pt;
    }

    table p:first-of-type strong{
        font-size: 14pt;
    }
    """
    with open(file, "rb") as docx_f:
        result = mammoth.convert_to_html(docx_f)
        html = result.value + custom_css
    html_file = file.replace('docx', 'html')
    with open(html_file, 'w') as html_f:
        html_f.write(html)

def main():
    # Otwórz plik z listą produktu, przejdź po każdym z nich i uzyskaj odpowiednie dane.
    # Dla każdego z plików konwertuj załącznik z pdf do docx i html
    # Zapisz listę produktów w plikach produkty.xml oraz produkty.csv
    # W razie błędów wypisz je do pliku "error_log.txt"

    possible_end = ("Referencyjna Wartość Spożycia", "Sugerowana porcja produktu (wartość)", "Warunki przechowywania", "Opakowanie oznakowane datą minimalnej trwałości", \
                        "Obowiązkowe dodatkowe informacje wymagane na etykiecie", "Kraj pochodzenia", "Sposób przygotowania", "Propozycja podania", "Przybliżona liczba porcji", \
                        "Zawartość netto", "Pozostałe informacje", "Nazwa opisowa produktu", "Nazwa firmy","Adres", "Forma kontaktu", "Inne informacje dla konsumenta"\
                        "Obecność alergenu", "Wartość odżywcza/ 100g", "kJ", "kcal", "Contains : ")

    unavailable_products = {
        'Indeks/SupplerItemCode': [],
        'Kod kreskowy/EAN': []
    }

    session, headers = get_bearer_token()
    first_run = True    
    try:
        print('Otwieram plik z listą produktów do wyszukania...')
        query_items = pd.read_excel('indeks+ean.xlsx')
    except:
        print('Błąd otwierania pliku!')
    else:
        print('Plik otwarty pomyślnie')
        
        for indx, row in query_items.iterrows():
            try:
                query_index = str(int(row['Indeks/SupplerItemCode']))
                query_ean = str(int(row['Kod kreskowy/EAN']))
                session, headers, product, unavailable_products = get_prd_data(session, headers, query_index, query_ean, unavailable_products)

                if product and product.attachment == True:
                    print('Otwieram plik pdf...')
                    ocr_data = Processed_OCR(product.product_name)
                    
                    pdf_file_name = 'zalaczniki_pdf\\' + product.ean + '.pdf'
                    split_text = open_pdf(pdf_file_name)
                    if not split_text:
                        break
                    split_text = rem_prd_id(split_text, product.product_name)
                    
                    split_text, ocr_data.allergens, ocr_data.add_descr = get_ean_allergens_add_discr(split_text)
                    split_text, ocr_data.ean = rem_ean(split_text, product.ean)
                    split_text, ocr_data.ingridients = get_ingridients(split_text)
                    split_text, ocr_data.allergens = get_allergens(split_text, ocr_data.allergens)
                    split_text, ocr_data.additional_info = get_additional_info(split_text)
                    split_text, ocr_data.before_table, ocr_data.table1, ocr_data.table2, ocr_data.table3 = get_tables(split_text)
                    if ocr_data.ingridients == [] and split_text != []:
                        for el in split_text:
                            if any(x in el for x in possible_end):
                                break
                            ocr_data.ingridients.append(el)
                    docx_file = write_to_docx(ocr_data)
                    write_to_html(docx_file)
                if product:
                    print('Dodaję produkt do bazy danych... ')
                    instock_products = pd.DataFrame(product.as_dict())
                    if first_run == True:
                        if path.exists('produkty.csv'):
                            products = pd.read_csv('produkty.csv')
                            if products.empty:
                                instock_products.to_csv('produkty.csv', index=False, mode='a', header=True)
                        else:
                            instock_products.to_csv('produkty.csv', index=False, mode='a', header=True)
                        first_run = False
                    else: 
                        instock_products.to_csv('produkty.csv', index=False, mode='a', header=False)
                    #instock_products.to_xml('produkty.xml')

                    soldout_products = pd.DataFrame(unavailable_products)
                    soldout_products.to_excel('niedostepne_produkty.xlsx', index=False)
                    print('Produkt zapisany.\n')
            except:
                with open('error_log.txt', 'a') as error_log:
                    error_log.write(query_index + ' ' + query_ean + '\n')
                    err = traceback.format_exc()
                    error_log.write(err)

    instock_products = pd.read_csv('produkty.csv')
    #instock_products.to_csv('produkty.csv', index=False)
    print('Zapisuję dostępne produkty do pliku "produkty.xml"...\n')
    instock_products.to_xml('produkty.xml')

    soldout_products = pd.DataFrame(unavailable_products)
    print('Zapisuję niedostepne produkty do pliku "niedostepne_produkty.xlsx"...\n')
    soldout_products.to_excel('niedostepne_produkty.xlsx', index=False)
    

if __name__=="__main__":
    main()